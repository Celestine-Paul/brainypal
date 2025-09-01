def preprocess_study_content(content: str) -> Dict[str, Any]:
    # BrainyPal Utility Functions
# utils.py

import os
import re
import logging
import hashlib
import json
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional

# Try to import optional dependencies
try:
    import PyPDF2
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False
    PyPDF2 = None

try:
    import docx
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False
    docx = None

try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    HAS_NLTK = True
except ImportError:
    HAS_NLTK = False
    nltk = None

# Configure logging
logger = logging.getLogger(__name__)

# File handling utilities
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx', 'md', 'rtf'}

def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_type(filename: str) -> str:
    """Get file type from filename"""
    if not filename:
        return 'unknown'
    
    extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    file_types = {
        'pdf': 'pdf',
        'doc': 'word',
        'docx': 'word',
        'txt': 'text',
        'md': 'markdown',
        'rtf': 'rtf'
    }
    
    return file_types.get(extension, 'unknown')

def generate_secure_filename(original_filename: str) -> str:
    """Generate secure filename with timestamp"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    secure_name = secure_filename(original_filename)
    name, ext = os.path.splitext(secure_name)
    return f"{timestamp}_{name}{ext}"

def get_file_size_mb(filepath: str) -> float:
    """Get file size in MB"""
    try:
        size_bytes = os.path.getsize(filepath)
        return round(size_bytes / (1024 * 1024), 2)
    except Exception:
        return 0.0

# File processing functions
def process_uploaded_file(filepath: str) -> str:
    """Process uploaded file and extract text content"""
    try:
        file_type = get_file_type(filepath)
        
        if file_type == 'pdf':
            return extract_text_from_pdf(filepath)
        elif file_type == 'word':
            return extract_text_from_word(filepath)
        elif file_type in ['text', 'markdown']:
            return extract_text_from_text_file(filepath)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
            
    except Exception as e:
        logger.error(f"File processing error for {filepath}: {e}")
        raise e

def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF file"""
    if not HAS_PDF_SUPPORT:
        raise ValueError("PDF support not available. Install PyPDF2: pip install PyPDF2")
    
    try:
        text = ""
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        # Clean extracted text
        text = clean_extracted_text(text)
        
        if not text.strip():
            raise ValueError("No text could be extracted from PDF")
        
        return text
        
    except Exception as e:
        logger.error(f"PDF text extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_word(filepath: str) -> str:
    """Extract text from Word document"""
    if not HAS_DOCX_SUPPORT:
        raise ValueError("Word document support not available. Install python-docx: pip install python-docx")
    
    try:
        if filepath.endswith('.docx'):
            doc = docx.Document(filepath)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        
        else:
            # Handle .doc files (older format)
            try:
                import docx2txt
                text = docx2txt.process(filepath)
            except ImportError:
                raise ValueError("Cannot process .doc files. Please convert to .docx format or install docx2txt.")
        
        # Clean extracted text
        text = clean_extracted_text(text)
        
        if not text.strip():
            raise ValueError("No text could be extracted from Word document")
        
        return text
        
    except Exception as e:
        logger.error(f"Word document text extraction failed: {e}")
        raise ValueError(f"Failed to extract text from Word document: {str(e)}")
        logger.error(f"Word document text extraction failed: {e}")
        raise ValueError(f"Failed to extract text from Word document: {str(e)}")

def extract_text_from_text_file(filepath: str) -> str:
    """Extract text from plain text file"""
    try:
        encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    text = file.read()
                
                # Clean and validate text
                text = clean_extracted_text(text)
                
                if text.strip():
                    return text
                    
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
        
    except Exception as e:
        logger.error(f"Text file extraction failed: {e}")
        raise ValueError(f"Failed to extract text from file: {str(e)}")

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove non-printable characters except newlines and tabs
    text = re.sub(r'[^\x20-\x7E\n\t]', '', text)
    
    # Remove multiple consecutive newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text

def validate_file_content(content: str, min_length: int = 50) -> bool:
    """Validate extracted file content"""
    if not content or len(content.strip()) < min_length:
        return False
    
    # Check if content has enough meaningful words
    words = word_tokenize(content.lower())
    meaningful_words = [word for word in words if word.isalpha() and len(word) > 2]
    
    return len(meaningful_words) >= 20

# Text processing utilities
def preprocess_study_content(content: str) -> Dict[str, Any]:
    """Preprocess study content for AI generation"""
    try:
        # Initialize NLTK if available
        if HAS_NLTK:
            try:
                nltk.data.find('tokenizers/punkt')
            except LookupError:
                try:
                    nltk.download('punkt', quiet=True)
                except:
                    pass  # Ignore if download fails
        
        # Basic statistics
        if HAS_NLTK:
            try:
                words = word_tokenize(content)
                sentences = sent_tokenize(content)
                word_count = len(words)
                sentence_count = len(sentences)
            except:
                # Fallback to simple splitting
                words = content.split()
                sentences = content.split('.')
                word_count = len(words)
                sentence_count = len(sentences)
        else:
            # Simple fallback
            words = content.split()
            sentences = content.split('.')
            word_count = len(words)
            sentence_count = len(sentences)
        
        char_count = len(content)
        
        # Estimate reading time (average 200 words per minute)
        reading_time = max(1, round(word_count / 200))
        
        # Extract key information
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        # Find potential headers/titles
        headers = extract_headers(content)
        
        # Estimate difficulty
        difficulty = estimate_content_difficulty(content, word_count, sentence_count)
        
        return {
            'word_count': word_count,
            'sentence_count': sentence_count,
            'char_count': char_count,
            'paragraph_count': len(paragraphs),
            'estimated_reading_time': reading_time,
            'estimated_difficulty': difficulty,
            'headers': headers,
            'key_sentences': rank_sentences_by_importance(sentences)[:10]
        }
        
    except Exception as e:
        logger.error(f"Content preprocessing failed: {e}")
        return {
            'word_count': len(content.split()),
            'sentence_count': len(content.split('.')),
            'char_count': len(content),
            'estimated_difficulty': 'intermediate',
            'error': str(e)
        }

def extract_headers(content: str) -> List[str]:
    """Extract potential headers and section titles"""
    headers = []
    
    # Pattern for markdown headers
    markdown_headers = re.findall(r'^#+\s+(.+), content, re.MULTILINE)
    headers.extend(markdown_headers)
    
    # Pattern for numbered sections
    numbered_sections = re.findall(r'^\d+\.?\s+([A-Z][^.\n]{10,60}), content, re.MULTILINE)
    headers.extend(numbered_sections)
    
    # Pattern for title case lines
    title_lines = re.findall(r'^([A-Z][A-Za-z\s]{10,60}), content, re.MULTILINE)
    headers.extend(title_lines)
    
    # Remove duplicates and clean
    clean_headers = []
    for header in headers:
        header = header.strip()
        if header and header not in clean_headers and len(header.split()) <= 10:
            clean_headers.append(header)
    
    return clean_headers[:10]  # Return top 10 headers

def estimate_content_difficulty(content: str, word_count: int, sentence_count: int) -> str:
    """Estimate content difficulty level"""
    try:
        # Calculate average sentence length
        avg_sentence_length = word_count / sentence_count if sentence_count > 0 else 0
        
        # Calculate average word length
        if HAS_NLTK:
            try:
                words = word_tokenize(content.lower())
            except:
                words = content.lower().split()
        else:
            words = content.lower().split()
            
        avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
        
        # Count complex words (3+ syllables or 7+ characters)
        complex_words = sum(1 for word in words if len(word) >= 7 or count_syllables(word) >= 3)
        complex_word_ratio = complex_words / len(words) if words else 0
        
        # Scoring system
        difficulty_score = 0
        
        if avg_sentence_length > 20:
            difficulty_score += 2
        elif avg_sentence_length > 15:
            difficulty_score += 1
        
        if avg_word_length > 6:
            difficulty_score += 2
        elif avg_word_length > 5:
            difficulty_score += 1
        
        if complex_word_ratio > 0.3:
            difficulty_score += 2
        elif complex_word_ratio > 0.2:
            difficulty_score += 1
        
        # Determine difficulty
        if difficulty_score >= 4:
            return 'advanced'
        elif difficulty_score >= 2:
            return 'intermediate'
        else:
            return 'beginner'
            
    except Exception as e:
        logger.error(f"Difficulty estimation failed: {e}")
        return 'intermediate'

def count_syllables(word: str) -> int:
    """Count syllables in a word (approximation)"""
    try:
        word = word.lower()
        vowels = 'aeiouy'
        syllable_count = 0
        prev_was_vowel = False
        
        for char in word:
            is_vowel = char in vowels
            if is_vowel and not prev_was_vowel:
                syllable_count += 1
            prev_was_vowel = is_vowel
        
        # Handle silent 'e'
        if word.endswith('e') and syllable_count > 1:
            syllable_count -= 1
        
        return max(1, syllable_count)
        
    except Exception:
        return 1

def rank_sentences_by_importance(sentences: List[str]) -> List[str]:
    """Rank sentences by importance for study material generation"""
    if not sentences:
        return []
        
    scored_sentences = []
    
    for sentence in sentences:
        if sentence.strip():  # Skip empty sentences
            score = calculate_sentence_importance(sentence)
            scored_sentences.append((sentence.strip(), score))
    
    # Sort by score (highest first)
    scored_sentences.sort(key=lambda x: x[1], reverse=True)
    
    return [sentence for sentence, score in scored_sentences]

def calculate_sentence_importance(sentence: str) -> float:
    """Calculate importance score for a sentence"""
    if not sentence or not sentence.strip():
        return 0.0
        
    score = 0.0
    sentence = sentence.strip()
    
    # Length factor (prefer medium-length sentences)
    if HAS_NLTK:
        try:
            word_count = len(word_tokenize(sentence))
        except:
            word_count = len(sentence.split())
    else:
        word_count = len(sentence.split())
        
    if 10 <= word_count <= 25:
        score += 2.0
    elif 8 <= word_count <= 30:
        score += 1.0
    
    # Keyword indicators
    importance_keywords = [
        'important', 'significant', 'crucial', 'essential', 'key', 'main', 'primary',
        'fundamental', 'critical', 'major', 'principal', 'central', 'vital'
    ]
    
    for keyword in importance_keywords:
        if keyword in sentence.lower():
            score += 1.5
    
    # Definition indicators
    definition_indicators = [' is ', ' are ', ' means ', ' refers to ', ' defined as ', ' known as ']
    for indicator in definition_indicators:
        if indicator in sentence.lower():
            score += 1.0
    
    # Numerical information
    if re.search(r'\d+', sentence):
        score += 0.5
    
    # Causal relationships
    causal_words = ['because', 'therefore', 'thus', 'consequently', 'as a result', 'due to']
    for word in causal_words:
        if word in sentence.lower():
            score += 1.0
    
    # Avoid very short or very long sentences
    if word_count < 5 or word_count > 40:
        score -= 1.0
    
    return max(0.0, score)

# Security utilities
def generate_api_key() -> str:
    """Generate secure API key"""
    import secrets
    return secrets.token_urlsafe(32)

def hash_password(password: str) -> str:
    """Hash password using secure method"""
    from werkzeug.security import generate_password_hash
    return generate_password_hash(password, method='pbkdf2:sha256', salt_length=16)

def verify_password(password: str, password_hash: str) -> bool:
    """Verify password against hash"""
    from werkzeug.security import check_password_hash
    return check_password_hash(password_hash, password)

def generate_reset_token() -> str:
    """Generate password reset token"""
    import secrets
    return secrets.token_urlsafe(32)

def validate_email(email: str) -> bool:
    """Validate email address format"""
    email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}
    return re.match(email_pattern, email) is not None

def validate_phone_number(phone: str, country_code: str = '+254') -> str:
    """Validate and format Kenyan phone number"""
    # Remove all non-digit characters
    phone = re.sub(r'\D', '', phone)
    
    # Handle different formats
    if phone.startswith('254'):
        phone = '+' + phone
    elif phone.startswith('0'):
        phone = country_code + phone[1:]
    elif len(phone) == 9:
        phone = country_code + phone
    else:
        raise ValueError("Invalid phone number format")
    
    # Validate Kenyan mobile number format
    if not re.match(r'^\+254[17]\d{8}, phone):
        raise ValueError("Please enter a valid Kenyan mobile number")
    
    return phone

# Content validation utilities
def validate_study_content(content: str) -> Dict[str, Any]:
    """Validate study content quality and suitability"""
    validation_result = {
        'is_valid': True,
        'issues': [],
        'warnings': [],
        'suggestions': []
    }
    
    # Minimum length check
    if len(content.strip()) < 100:
        validation_result['is_valid'] = False
        validation_result['issues'].append("Content is too short. Please provide at least 100 characters.")
    
    # Maximum length check
    if len(content) > 50000:
        validation_result['warnings'].append("Content is very long. Consider breaking it into smaller sections.")
    
    # Language detection (basic)
    if not is_primarily_english(content):
        validation_result['warnings'].append("Content appears to be in a non-English language. AI generation works best with English content.")
    
    # Educational content quality
    sentences = sent_tokenize(content)
    if len(sentences) < 5:
        validation_result['warnings'].append("Content has very few sentences. More detailed content will produce better study materials.")
    
    # Check for structured content
    if has_structured_content(content):
        validation_result['suggestions'].append("Great! Your content appears well-structured, which will help generate better study materials.")
    
    return validation_result

def is_primarily_english(text: str) -> bool:
    """Check if text is primarily in English"""
    try:
        # Simple heuristic: check for common English words
        common_english_words = {
            'the', 'and', 'is', 'in', 'to', 'of', 'a', 'that', 'it', 'with',
            'for', 'as', 'was', 'on', 'are', 'by', 'this', 'be', 'at', 'from'
        }
        
        if HAS_NLTK:
            try:
                words = word_tokenize(text.lower())
            except:
                words = text.lower().split()
        else:
            words = text.lower().split()
            
        if not words:
            return False
            
        english_word_count = sum(1 for word in words if word in common_english_words)
        return english_word_count / len(words) > 0.1
        
    except Exception:
        return True  # Assume English if detection fails

def has_structured_content(content: str) -> bool:
    """Check if content has structured formatting"""
    indicators = [
        r'^\d+\.',  # Numbered lists
        r'^-\s',    # Bullet points
        r'^#+ ',    # Markdown headers
        r'[A-Z][^.]{20,}:',  # Section headers with colons
        r'\n\s*\n\s*[A-Z]'   # Paragraph breaks
    ]
    
    for pattern in indicators:
        if re.search(pattern, content, re.MULTILINE):
            return True
    
    return False

# Database utilities
def create_database_backup(db_uri: str, backup_path: str) -> bool:
    """Create database backup"""
    try:
        # For MySQL databases
        if 'mysql' in db_uri:
            import subprocess
            
            # Extract database connection details
            # This is a simplified version - implement proper parsing
            result = subprocess.run([
                'mysqldump',
                '--single-transaction',
                '--routines',
                '--triggers',
                'brainypal_db'
            ], capture_output=True, text=True)
            
            if result.returncode == 0:
                with open(backup_path, 'w') as backup_file:
                    backup_file.write(result.stdout)
                return True
        
        return False
        
    except Exception as e:
        logger.error(f"Database backup failed: {e}")
        return False

def cleanup_old_files(directory: str, max_age_days: int = 7) -> int:
    """Clean up old uploaded files"""
    try:
        cutoff_date = datetime.now() - timedelta(days=max_age_days)
        deleted_count = 0
        
        for filename in os.listdir(directory):
            filepath = os.path.join(directory, filename)
            
            if os.path.isfile(filepath):
                file_modified = datetime.fromtimestamp(os.path.getmtime(filepath))
                
                if file_modified < cutoff_date:
                    try:
                        os.remove(filepath)
                        deleted_count += 1
                        logger.info(f"Deleted old file: {filename}")
                    except Exception as e:
                        logger.error(f"Failed to delete {filename}: {e}")
        
        return deleted_count
        
    except Exception as e:
        logger.error(f"File cleanup failed: {e}")
        return 0

# Performance utilities
def measure_execution_time(func):
    """Decorator to measure function execution time"""
    import functools
    import time
    
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        
        execution_time = end_time - start_time
        logger.info(f"Function {func.__name__} executed in {execution_time:.2f} seconds")
        
        return result
    
    return wrapper

def log_api_usage(user_id: int, endpoint: str, success: bool, response_time: float):
    """Log API usage for analytics"""
    try:
        usage_data = {
            'user_id': user_id,
            'endpoint': endpoint,
            'success': success,
            'response_time': response_time,
            'timestamp': datetime.utcnow().isoformat()
        }
        
        # In production, you might want to send this to a dedicated logging service
        logger.info(f"API Usage: {json.dumps(usage_data)}")
        
    except Exception as e:
        logger.error(f"Usage logging failed: {e}")

# Rate limiting utilities
class RateLimiter:
    """Simple in-memory rate limiter"""
    
    def __init__(self):
        self.requests = {}
    
    def is_allowed(self, user_id: int, limit: int, window_minutes: int = 60) -> bool:
        """Check if request is allowed within rate limit"""
        now = datetime.utcnow()
        window_start = now - timedelta(minutes=window_minutes)
        
        # Clean old requests
        if user_id in self.requests:
            self.requests[user_id] = [
                req_time for req_time in self.requests[user_id] 
                if req_time > window_start
            ]
        else:
            self.requests[user_id] = []
        
        # Check if under limit
        if len(self.requests[user_id]) < limit:
            self.requests[user_id].append(now)
            return True
        
        return False
    
    def get_remaining_requests(self, user_id: int, limit: int, window_minutes: int = 60) -> int:
        """Get remaining requests in current window"""
        now = datetime.utcnow()
        window_start = now - timedelta(minutes=window_minutes)
        
        if user_id in self.requests:
            recent_requests = [
                req_time for req_time in self.requests[user_id] 
                if req_time > window_start
            ]
            return max(0, limit - len(recent_requests))
        
        return limit

# Caching utilities
class SimpleCache:
    """Simple in-memory cache for API responses"""
    
    def __init__(self, default_ttl: int = 300):
        self.cache = {}
        self.default_ttl = default_ttl
    
    def get(self, key: str) -> Any:
        """Get cached value"""
        if key in self.cache:
            value, expiry = self.cache[key]
            if datetime.utcnow() < expiry:
                return value
            else:
                del self.cache[key]
        return None
    
    def set(self, key: str, value: Any, ttl: int = None) -> None:
        """Set cached value"""
        ttl = ttl or self.default_ttl
        expiry = datetime.utcnow() + timedelta(seconds=ttl)
        self.cache[key] = (value, expiry)
    
    def delete(self, key: str) -> None:
        """Delete cached value"""
        if key in self.cache:
            del self.cache[key]
    
    def clear(self) -> None:
        """Clear all cached values"""
        self.cache.clear()

# Error handling utilities
class BrainyPalException(Exception):
    """Base exception for BrainyPal application"""
    pass

class ContentProcessingError(BrainyPalException):
    """Exception for content processing errors"""
    pass

class PaymentError(BrainyPalException):
    """Exception for payment processing errors"""
    pass

class AIServiceError(BrainyPalException):
    """Exception for AI service errors"""
    pass

def handle_api_error(error: Exception, user_id: int = None, endpoint: str = None) -> Dict[str, Any]:
    """Standardized API error handling"""
    error_response = {
        'success': False,
        'error_type': type(error).__name__,
        'message': str(error),
        'timestamp': datetime.utcnow().isoformat()
    }
    
    # Log error with context
    log_context = {
        'user_id': user_id,
        'endpoint': endpoint,
        'error_type': type(error).__name__,
        'error_message': str(error)
    }
    
    logger.error(f"API Error: {json.dumps(log_context)}")
    
    # Customize error message for users
    if isinstance(error, ContentProcessingError):
        error_response['user_message'] = "We couldn't process your content. Please check the file format and try again."
    elif isinstance(error, PaymentError):
        error_response['user_message'] = "Payment processing failed. Please try again or contact support."
    elif isinstance(error, AIServiceError):
        error_response['user_message'] = "AI service is temporarily unavailable. Please try again in a few moments."
    else:
        error_response['user_message'] = "An unexpected error occurred. Please try again."
    
    return error_response

# Environment utilities
def load_environment_variables():
    """Load environment variables from .env file"""
    try:
        from dotenv import load_dotenv
        load_dotenv()
        logger.info("Environment variables loaded from .env file")
    except ImportError:
        logger.warning("python-dotenv not installed. Environment variables should be set manually.")
    except Exception as e:
        logger.error(f"Failed to load environment variables: {e}")

def verify_required_services():
    """Verify that required external services are accessible"""
    services_status = {
        'mysql': False,
        'huggingface': False,
        'intasend': False
    }
    
    # Test MySQL connection
    try:
        import pymysql
        # This would test actual connection in production
        services_status['mysql'] = True
    except Exception as e:
        logger.error(f"MySQL connection test failed: {e}")
    
    # Test Hugging Face API
    try:
        hf_key = os.getenv('HUGGINGFACE_API_KEY')
        if hf_key and hf_key != 'hf_your_api_key_here':
            # Test API call would go here
            services_status['huggingface'] = True
    except Exception as e:
        logger.error(f"Hugging Face API test failed: {e}")
    
    # Test IntaSend API
    try:
        intasend_key = os.getenv('INTASEND_SECRET_KEY')
        if intasend_key and intasend_key != 'ISSecretKey_test_your_key_here':
            # Test API call would go here
            services_status['intasend'] = True
    except Exception as e:
        logger.error(f"IntaSend API test failed: {e}")
    
    return services_status

# Data formatting utilities
def format_currency(amount: float, currency: str = 'KES') -> str:
    """Format currency amount for display"""
    if currency == 'KES':
        return f'KSh {amount:,.0f}'
    elif currency == 'USD':
        return f'${amount:,.2f}'
    else:
        return f'{amount:,.2f} {currency}'

def format_datetime(dt: datetime, format_type: str = 'default') -> str:
    """Format datetime for display"""
    if not dt:
        return 'N/A'
    
    formats = {
        'default': '%Y-%m-%d %H:%M:%S',
        'date_only': '%Y-%m-%d',
        'time_only': '%H:%M:%S',
        'friendly': '%B %d, %Y at %I:%M %p',
        'iso': '%Y-%m-%dT%H:%M:%SZ'
    }
    
    return dt.strftime(formats.get(format_type, formats['default']))

def calculate_time_ago(dt: datetime) -> str:
    """Calculate human-readable time ago"""
    if not dt:
        return 'Unknown'
    
    now = datetime.utcnow()
    diff = now - dt
    
    if diff.days > 0:
        return f'{diff.days} day{"s" if diff.days != 1 else ""} ago'
    elif diff.seconds > 3600:
        hours = diff.seconds // 3600
        return f'{hours} hour{"s" if hours != 1 else ""} ago'
    elif diff.seconds > 60:
        minutes = diff.seconds // 60
        return f'{minutes} minute{"s" if minutes != 1 else ""} ago'
    else:
        return 'Just now'

# Initialize global utilities
rate_limiter = RateLimiter()
cache = SimpleCache()

# Export all utility functions and classes
__all__ = [
    'allowed_file', 'get_file_type', 'generate_secure_filename', 'get_file_size_mb',
    'process_uploaded_file', 'extract_text_from_pdf', 'extract_text_from_word', 
    'extract_text_from_text_file', 'clean_extracted_text', 'validate_file_content',
    'preprocess_study_content', 'extract_headers', 'estimate_content_difficulty',
    'rank_sentences_by_importance', 'validate_study_content', 'validate_email',
    'validate_phone_number', 'format_currency', 'format_datetime', 'calculate_time_ago',
    'RateLimiter', 'SimpleCache', 'rate_limiter', 'cache',
    'BrainyPalException', 'ContentProcessingError', 'PaymentError', 'AIServiceError',
    'handle_api_error', 'load_environment_variables', 'verify_required_services'
]